# app.py — Psychiatric Dropout Risk (Monotone + Calibration + Pre/Post + Ablation + Capacity + Fairness + Follow-up window)
# 亮點：
# - 單調性 XGBoost +（可用則）isotonic calibration
# - Pre/Post 模式：模型避免用已安排追蹤(防洩漏)，但 Overlay 在兩種模式皆視追蹤為保護
# - Overlay（文獻/規則）可與模型混合（Blend），並在 Validation 做 Model/Overlay/Final 三軌 Ablation
# - Decision Curve、門檻可調、容量/工時計（社會資源治理）
# - 子群公平（性別/年齡段/主診斷）AUC + ECE
# - 單例頁 SHAP（model）+ Policy drivers 對齊卡（衝突標示⚠️）
# - What-if 面板（followups / compliance）
# - Null-safe + Vignettes 匯出
# - 新增：Follow-up window (days) + Chief problem 一致化 + LOS/診斷 sanity check

import os, re, math
import streamlit as st
import pandas as pd
import numpy as np
import joblib
import shap
import matplotlib.pyplot as plt
from io import BytesIO

# 可選：Word 匯出 SOP
try:
    from docx import Document
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

st.set_page_config(page_title="Psychiatric Dropout Risk", layout="wide")
st.title("🧠 Psychiatric Dropout Risk Predictor (Monotone + Calibrated)")

# ==== Math helpers ====
def _sigmoid(x): return 1.0 / (1.0 + np.exp(-x))
def _logit(p, eps=1e-6):
    p = float(np.clip(p, eps, 1 - eps)); return np.log(p / (1 - p))
def _logit_vec(p, eps=1e-6):
    p = np.clip(p, eps, 1 - eps); return np.log(p / (1 - p))

# ==== Global knobs ====
CAL_LOGIT_SHIFT = float(os.getenv("RISK_CAL_SHIFT", "0.0"))
BORDER_BAND = 7
BLEND_W_DEFAULT = 0.30        # Final = (1-BLEND)*Model + BLEND*Overlay
SOFT_UPLIFT = {"floor": 0.60, "add": 0.10, "cap": 0.90}  # 自傷保護

# Overlay safety controls
OVERLAY_SCALE = 0.75
DELTA_CLIP   = 1.00
TEMP         = 1.20

# ====== Options & schema ======
DIAG_LIST = [
    "Schizophrenia","Bipolar","Depression","Personality Disorder",
    "Substance Use Disorder","Dementia","Anxiety","PTSD","OCD","ADHD","Other/Unknown"
]
BIN_YESNO = ["Yes","No"]
GENDER_LIST = ["Male","Female","Other/Unknown"]

TEMPLATE_COLUMNS = [
    "age","length_of_stay","num_previous_admissions",
    "medication_compliance_score","family_support_score","post_discharge_followups",
    "gender_Male","gender_Female","gender_Other/Unknown",
] + [f"diagnosis_{d}" for d in DIAG_LIST] + [
    "has_recent_self_harm_Yes","has_recent_self_harm_No",
    "self_harm_during_admission_Yes","self_harm_during_admission_No",
]

# ====== Null-safe defaults ======
DEFAULTS = {
    "age": 40.0,
    "length_of_stay": 21.0,  # 調高為更貼近台灣急性病房
    "num_previous_admissions": 0.0,
    "medication_compliance_score": 5.0,
    "family_support_score": 5.0,
    "post_discharge_followups": 0.0,
}
NUMERIC_KEYS = list(DEFAULTS.keys())

def _num_or_default(x, key):
    try: v = float(x)
    except Exception: v = np.nan
    if pd.isna(v): v = DEFAULTS.get(key, 0.0)
    return v

def fill_defaults_single_row(X1: pd.DataFrame):
    i = X1.index[0]
    for k in NUMERIC_KEYS:
        if k in X1.columns:
            X1.at[i, k] = _num_or_default(X1.at[i, k], k)
    diag_cols = [c for c in X1.columns if c.startswith("diagnosis_")]
    if diag_cols and (X1.loc[i, diag_cols].sum() == 0):
        col = "diagnosis_Other/Unknown"
        if col in X1.columns: X1.at[i, col] = 1
    oh_cols = [c for c in X1.columns if "_" in c and c not in NUMERIC_KEYS]
    if oh_cols: X1.loc[i, oh_cols] = X1.loc[i, oh_cols].fillna(0)

def fill_defaults_batch(df_feat: pd.DataFrame):
    for k in NUMERIC_KEYS:
        if k in df_feat.columns:
            df_feat[k] = pd.to_numeric(df_feat[k], errors="coerce").fillna(DEFAULTS[k])
    oh_cols = [c for c in df_feat.columns if "_" in c and c not in NUMERIC_KEYS]
    if oh_cols: df_feat[oh_cols] = df_feat[oh_cols].fillna(0)
    diag_cols = [c for c in df_feat.columns if c.startswith("diagnosis_")]
    if diag_cols:
        none_diag = (df_feat[diag_cols].sum(axis=1) == 0)
        if none_diag.any() and "diagnosis_Other/Unknown" in df_feat.columns:
            df_feat.loc[none_diag, "diagnosis_Other/Unknown"] = 1

# ====== Overlay policy（log-odds, 文獻啟發，方向單調）======
POLICY = {
    "per_prev_admission": 0.18,       # ↑ 住院史
    "per_point_low_support": 0.20,    # 家庭支持低
    "per_followup": -0.18,            # 追蹤愈多愈保護（之後乘以視窗權重）
    "no_followup_extra": 0.30,        # 沒追蹤加罰
    "los_short": 0.40, "los_mid": 0.00, "los_mid_high": 0.15, "los_long": 0.30,
    "age_young": 0.10, "age_old": 0.10,
    "diag": {
        "Personality Disorder":    0.35,
        "Substance Use Disorder":  0.35,
        "Bipolar":                 0.10,
        "PTSD":                    0.10,
        "Schizophrenia":           0.10,
        "Depression":              0.05,
        "Anxiety":                 0.00,
        "OCD":                     0.00,
        "Dementia":                0.00,
        "ADHD":                    0.00,
        "Other/Unknown":           0.00,
    },
    "x_sud_lowcomp": 0.30,     # SUD × 低順從
    "x_pd_shortlos": 0.10,     # PD × 短住院
    # compliance：低→加分，高→保護
    "per_point_low_compliance": 0.22,
    "per_point_high_compliance_protect": -0.06,
}

# ====== UI — Sidebar ======
with st.sidebar:
    st.header("Patient Info")

    # Chief problem（與自傷旗標做一致化）
    complaint = st.selectbox(
        "Chief problem",
        [
            "Routine follow-up",
            "Medication issue",
            "Psychotic exacerbation (agitation/violence)",
            "Severe depressive symptoms / functional decline",
            "Suicidal ideation / Self-harm risk",
            "Unable self-care / caregiver overload",
            "Medication side-effects",
            "Diagnostic clarification (complex psychopathology)",
            "Other/Unknown",
        ],
        index=0
    )

    age = st.slider("Age", 18, 95, 35)
    gender = st.selectbox("Gender", GENDER_LIST, index=0)
    diagnoses = st.multiselect("Diagnoses (multi-select)", DIAG_LIST, default=[])
    # 調高預設 LOS
    length_of_stay = st.slider("Length of Stay (days)", 1, 90, 21)
    num_adm = st.slider("Previous Admissions (1y)", 0, 15, 1)
    compliance = st.slider("Medication Compliance Score (0–10)", 0.0, 10.0, 5.0)
    support = st.slider("Family Support Score (0–10)", 0.0, 10.0, 5.0)

    # 追蹤次數 + 視窗天數
    followups = st.slider("Post-discharge Followups (count)", 0, 10, 2)
    followup_window_days = st.number_input("Follow-up window (days)", min_value=3, max_value=60, value=14, step=1)

    recent_self_harm = st.radio("Recent Self-harm", BIN_YESNO, index=1)
    selfharm_adm = st.radio("Self-harm During Admission", BIN_YESNO, index=1)

    mode = st.radio("Mode", ["Pre-planning (no followup as feature)", "Post-planning (monitoring)"], index=0)
    use_followups_feature = (mode.startswith("Post"))

    st.markdown("---")
    with st.expander("Advanced (calibration & overlay)", expanded=False):
        cal_shift = st.slider("Global calibration (log-odds shift)", -1.0, 1.0, CAL_LOGIT_SHIFT, 0.05)
        blend_w  = st.slider("Blend weight (Final = (1-BLEND)*Model + BLEND*Overlay)", 0.0, 1.0, BLEND_W_DEFAULT, 0.05)
        overlay_scale = st.slider("Overlay scale", 0.0, 1.0, OVERLAY_SCALE, 0.05)
        delta_clip = st.slider("Overlay delta clip |log-odds|", 0.0, 2.0, DELTA_CLIP, 0.05)
        temp_val = st.slider("Temperature (>1 softer probs)", 0.5, 3.0, TEMP, 0.05)

    CAL_LOGIT_SHIFT = cal_shift
    BLEND_W = blend_w
    OVERLAY_SCALE = overlay_scale
    DELTA_CLIP = delta_clip
    TEMP = temp_val

# ====== Helpers (one-hot, alignment) ======
def set_onehot_by_prefix(df, prefix, value):
    col = f"{prefix}_{value}"
    if col in df.columns: df.at[0, col] = 1

def set_onehot_by_prefix_multi(df, prefix, values):
    for v in values:
        col = f"{prefix}_{v}"
        if col in df.columns: df.at[0, col] = 1

def flag_yes(row, prefix):
    col = f"{prefix}_Yes"; return (col in row.index) and (row[col] == 1)

def risk_bins(score, mod=20, high=40, band=BORDER_BAND):
    if score >= high + band: return "High"
    if score >= high - band: return "Moderate–High"
    if score >= mod + band:  return "Moderate"
    if score >= mod - band:  return "Low–Moderate"
    return "Low"

def proba_to_percent(p): return float(p) * 100
def proba_to_score(p): return int(round(proba_to_percent(p)))

# ====== Model load / train (monotone) + optional isotonic calibration ======
def get_monotone_constraints(feature_names):
    # +1: increasing risk; -1: decreasing risk; 0: no constraint
    mono_map = {
        "num_previous_admissions": +1,
        "medication_compliance_score": -1,
        "family_support_score": -1,
        "post_discharge_followups": -1,
        "length_of_stay": +1,
    }
    cons = [str(mono_map.get(f, 0)) for f in feature_names]
    return "(" + ",".join(cons) + ")"

def xgb_model_with_monotone(feature_names):
    import xgboost as xgb
    return xgb.XGBClassifier(
        n_estimators=500, max_depth=4, learning_rate=0.06,
        subsample=0.9, colsample_bytree=0.9, reg_lambda=1.0,
        random_state=42, tree_method="hist",
        objective="binary:logistic", eval_metric="logloss",
        monotone_constraints=get_monotone_constraints(feature_names)
    )

def try_load_model(path="dropout_model.pkl"):
    if not os.path.exists(path): return None
    try:
        mdl = joblib.load(path)
        if isinstance(mdl, dict) and "model" in mdl: return mdl["model"]
        return mdl
    except Exception:
        return None

def align_df_to_model(df: pd.DataFrame, m):
    names = None
    try:
        booster = getattr(m, "get_booster", lambda: None)()
        if booster is not None:
            nm = getattr(booster, "feature_names", None)
            if nm: names = list(nm)
    except Exception:
        pass
    if names:
        aligned = pd.DataFrame(0, index=df.index, columns=names, dtype=np.float32)
        inter = [c for c in names if c in df.columns]
        aligned.loc[:, inter] = df[inter].astype(np.float32).values
        return aligned, names
    # fallback
    out = df.astype(np.float32)
    return out, list(out.columns)

def train_demo_model_and_calibrator(columns):
    # 更寫實的合成資料（LOS 與診斷有偏移；SUD/PTSD/PD 有自傷較多）
    rng = np.random.default_rng(42)
    n = 12000
    X = pd.DataFrame(0, index=range(n), columns=columns, dtype=np.float32)
    X["age"] = rng.integers(16, 85, n)

    # base LOS
    los = rng.normal(21.0, 7.0, n)
    los = np.clip(los, 1, 60)

    # gender
    idx_gender = rng.integers(0, len(GENDER_LIST), n)
    for i, g in enumerate(GENDER_LIST): X.loc[idx_gender == i, f"gender_{g}"] = 1

    # 主診斷 + 共病
    idx_primary = rng.integers(0, len(DIAG_LIST), n)
    for i, d in enumerate(DIAG_LIST): X.loc[idx_primary == i, f"diagnosis_{d}"] = 1
    extra_probs = {"Substance Use Disorder": 0.20, "Anxiety": 0.20, "Depression": 0.25, "PTSD": 0.10}
    for d, pr in extra_probs.items(): X.loc[rng.random(n) < pr, f"diagnosis_{d}"] = 1

    # LOS 診斷偏移（SCZ/Bipolar 較長）
    scz = (X["diagnosis_Schizophrenia"] == 1).to_numpy()
    bp  = (X["diagnosis_Bipolar"] == 1).to_numpy()
    los = los + 7.0*scz + 4.0*bp
    X["length_of_stay"] = np.clip(los, 1, 90)

    X["num_previous_admissions"] = rng.poisson(0.9, n).clip(0, 12)
    X["medication_compliance_score"] = rng.normal(6.0, 2.5, n).clip(0, 10)
    X["family_support_score"] = rng.normal(5.0, 2.5, n).clip(0, 10)
    # followups（14 天窗內次數）
    X["post_discharge_followups"] = rng.integers(0, 6, n)

    # 自傷旗標（PD/PTSD/SUD 機率較高）
    base_r = rng.random(n)
    pd_flag   = (X["diagnosis_Personality Disorder"] == 1).to_numpy()
    ptsd_flag = (X["diagnosis_PTSD"] == 1).to_numpy()
    sud_flag  = (X["diagnosis_Substance Use Disorder"] == 1).to_numpy()
    prob_sh = 0.18 + 0.12*pd_flag + 0.10*ptsd_flag + 0.10*sud_flag
    r1 = (base_r < np.clip(prob_sh, 0, 0.9)).astype(int)
    r2 = (rng.random(n) < 0.12).astype(int)
    X.loc[r1 == 1, "has_recent_self_harm_Yes"] = 1; X.loc[r1 == 0, "has_recent_self_harm_No"] = 1
    X.loc[r2 == 1, "self_harm_during_admission_Yes"] = 1; X.loc[r2 == 0, "self_harm_during_admission_No"] = 1

    # 目標（與臨床一致的權重）
    beta0 = -0.60
    beta = {
        "has_recent_self_harm_Yes": 0.80,
        "self_harm_during_admission_Yes": 0.60,
        "prev_adm_ge2": 0.60,
        "medication_compliance_per_point": -0.25,
        "family_support_per_point": -0.20,
        "followups_per_visit": -0.15,         # 14d 視窗內每次觸點
        "length_of_stay_per_day": 0.04,       # 調整為 0.04
    }
    beta_diag = {
        "Personality Disorder": 0.35, "Substance Use Disorder": 0.35, "Bipolar": 0.10,
        "PTSD": 0.10, "Schizophrenia": 0.10, "Depression": 0.05,
        "Anxiety": 0.00, "OCD": 0.00, "Dementia": 0.00, "ADHD": 0.00, "Other/Unknown": 0.00,
    }
    prev_ge2 = (X["num_previous_admissions"] >= 2).astype(np.float32)
    logit = (beta0
             + beta["has_recent_self_harm_Yes"]       * X["has_recent_self_harm_Yes"]
             + beta["self_harm_during_admission_Yes"] * X["self_harm_during_admission_Yes"]
             + beta["prev_adm_ge2"]                   * prev_ge2
             + beta["medication_compliance_per_point"]* X["medication_compliance_score"]
             + beta["family_support_per_point"]       * X["family_support_score"]
             + beta["followups_per_visit"]            * X["post_discharge_followups"]
             + beta["length_of_stay_per_day"]         * X["length_of_stay"])
    for d, w in beta_diag.items(): logit = logit + w * X[f"diagnosis_{d}"]
    noise = rng.normal(0.0, 0.35, n).astype(np.float32)
    p = 1.0 / (1.0 + np.exp(-(logit + noise)))
    y = (rng.random(n) < p).astype(np.int32)

    # 單調性 XGB
    model = xgb_model_with_monotone(list(X.columns))
    model.fit(X, y)

    # Optional isotonic calibration
    calibrator = None
    try:
        from sklearn.model_selection import train_test_split
        from sklearn.calibration import CalibratedClassifierCV
        X_tr, X_ca, y_tr, y_ca = train_test_split(X, y, test_size=0.2, random_state=777)
        calibrator = CalibratedClassifierCV(model, cv="prefit", method="isotonic")
        calibrator.fit(X_ca, y_ca)
    except Exception:
        calibrator = None

    return model, calibrator, "demo (monotone + isotonic if available)"

def get_feat_names(m):
    try:
        b = m.get_booster()
        if getattr(b, "feature_names", None): return list(b.feature_names)
    except Exception: pass
    if hasattr(m, "feature_names_in_"): return list(m.feature_names_in_)
    return None

_loaded = try_load_model()
if _loaded is None:
    model, calibrator, model_source = train_demo_model_and_calibrator(TEMPLATE_COLUMNS)
else:
    model, calibrator, model_source = _loaded, None, "loaded from dropout_model.pkl"

def predict_model_proba(df_aligned: pd.DataFrame):
    probs = model.predict_proba(df_aligned, validate_features=False)[:, 1]
    if calibrator is not None:
        try:
            probs = calibrator.predict_proba(df_aligned)[:, 1]
        except Exception:
            pass
    return probs

# ====== Overlay（單例 + 驅動因子；含 follow-up window 權重）=====
def overlay_single_and_drivers(X1: pd.DataFrame, followup_window_days: int = 14):
    """
    Policy overlay on top of model logit; 不論 Pre/Post，Overlay 都把追蹤視為保護。
    followup_window_days：越短 -> 權重越大（以 14 天為基準）。
    """
    row = X1.iloc[0]
    drivers = []
    def add(label, val):
        if val != 0: drivers.append((label, float(val)))
        return val

    base_logit = _logit( float(model.predict_proba(X1, validate_features=False)[:,1][0]) )
    lz = base_logit

    adm = _num_or_default(row["num_previous_admissions"], "num_previous_admissions")
    comp = _num_or_default(row["medication_compliance_score"], "medication_compliance_score")
    sup  = _num_or_default(row["family_support_score"], "family_support_score")
    fup  = _num_or_default(row["post_discharge_followups"], "post_discharge_followups")
    los  = _num_or_default(row["length_of_stay"], "length_of_stay")
    agev = _num_or_default(row["age"], "age")

    # baseline drivers
    lz += add("More previous admissions", POLICY["per_prev_admission"] * min(int(adm), 5))
    lz += add("Low family support", POLICY["per_point_low_support"] * max(0.0, 5.0 - sup))
    lz += add("Low medication compliance", POLICY["per_point_low_compliance"] * max(0.0, 5.0 - comp))
    if comp >= 8:
        lz += add("High compliance (protective)", POLICY["per_point_high_compliance_protect"] * (comp - 7.0))

    # follow-ups with window weight
    wnd = max(3, min(60, int(followup_window_days)))
    w = np.sqrt(14.0 / wnd)              # window 越短 -> w 越大
    w = float(np.clip(w, 0.5, 1.5))
    eff_followups = fup * w
    lz += add(f"Follow-ups within {wnd}d (protective ×{w:.2f})", POLICY["per_followup"] * eff_followups)
    if fup == 0:
        lz += add("No follow-up scheduled", POLICY["no_followup_extra"])

    # LOS bucketing
    if los < 3: lz += add("Very short stay (<3d)", POLICY["los_short"])
    elif los <= 14: lz += add("Typical stay (3–14d)", POLICY["los_mid"])
    elif los <= 21: lz += add("Longish stay (15–21d)", POLICY["los_mid_high"])
    else: lz += add("Very long stay (>21d)", POLICY["los_long"])

    # age
    if agev < 21: lz += add("Young age (<21)", POLICY["age_young"])
    elif agev >= 75: lz += add("Older age (≥75)", POLICY["age_old"])

    # diagnoses
    for dx, wdx in POLICY["diag"].items():
        if X1.at[0, f"diagnosis_{dx}"] == 1:
            lz += add(f"Diagnosis: {dx}", wdx)

    # interactions
    if (X1.at[0, "diagnosis_Substance Use Disorder"] == 1) and (comp <= 3):
        lz += add("SUD × very low compliance", POLICY["x_sud_lowcomp"])
    if (X1.at[0, "diagnosis_Personality Disorder"] == 1) and (los < 3):
        lz += add("PD × very short stay", POLICY["x_pd_shortlos"])

    # scale + clip + calibration + temp
    delta = np.clip(OVERLAY_SCALE * (lz - base_logit), -DELTA_CLIP, DELTA_CLIP)
    lz2 = base_logit + delta + CAL_LOGIT_SHIFT
    p_overlay = _sigmoid(lz2 / TEMP)
    return float(p_overlay), drivers

# ====== Build single-row DF ======
X_single = pd.DataFrame(0, index=[0], columns=TEMPLATE_COLUMNS, dtype=float)
for k, v in {
    "age": age, "length_of_stay": float(length_of_stay), "num_previous_admissions": int(num_adm),
    "medication_compliance_score": float(compliance), "family_support_score": float(support),
    "post_discharge_followups": int(followups),
}.items(): X_single.at[0, k] = v
set_onehot_by_prefix(X_single, "gender", gender)
set_onehot_by_prefix_multi(X_single, "diagnosis", diagnoses)
set_onehot_by_prefix(X_single, "has_recent_self_harm", recent_self_harm)
set_onehot_by_prefix(X_single, "self_harm_during_admission", selfharm_adm)
fill_defaults_single_row(X_single)

# Chief problem 與自傷旗標一致化：若主訴含自傷但旗標皆 No → 自動補 recent_self_harm=Yes
if ("Suicidal" in complaint or "Self-harm" in complaint):
    if (X_single.at[0,"has_recent_self_harm_Yes"]==0) and (X_single.at[0,"self_harm_during_admission_Yes"]==0):
        X_single.at[0,"has_recent_self_harm_Yes"] = 1
        X_single.at[0,"has_recent_self_harm_No"]  = 0

# Pre-planning：模型不吃 followups（防洩漏）；Post 才保留（但 Overlay 兩者皆吃）
X_used = X_single.copy()
if not use_followups_feature:
    X_used.at[0, "post_discharge_followups"] = 0

# 預測（model + overlay + blend + uplift）
X_align, _ = align_df_to_model(X_used, model)
p_model = float(predict_model_proba(X_align)[0])
# Overlay 一律吃 followups（用視窗權重），即便 Pre-planning
p_overlay, drivers = overlay_single_and_drivers(X_single, followup_window_days=int(followup_window_days))
p_final = (1.0 - BLEND_W_DEFAULT) * p_model + BLEND_W_DEFAULT * p_overlay

# 自傷 uplift
if flag_yes(X_single.iloc[0], "has_recent_self_harm") or flag_yes(X_single.iloc[0], "self_harm_during_admission"):
    p_final = min(max(p_final, SOFT_UPLIFT["floor"]) + SOFT_UPLIFT["add"], SOFT_UPLIFT["cap"])

percent_model = proba_to_percent(p_model)
percent_overlay = proba_to_percent(p_overlay)
percent_final = proba_to_percent(p_final)
score = proba_to_score(p_final)
level = risk_bins(score)

# ====== Guards / input reasonableness ======
warns = []
# LOS 合理性
if length_of_stay > 60:
    warns.append("Length of stay > 60d is unusual; check data.")
if (num_adm > 10):
    warns.append("Previous admissions > 10 in 1y is uncommon; confirm definition.")
# PD/PTSD/SUD 常見自傷，但兩旗標都 No
_has_pd = X_single.at[0, "diagnosis_Personality Disorder"] == 1
_has_ptsd = X_single.at[0, "diagnosis_PTSD"] == 1
_has_sud = X_single.at[0, "diagnosis_Substance Use Disorder"] == 1
if (_has_pd or _has_ptsd or _has_sud) and (X_single.at[0,"has_recent_self_harm_Yes"]==0) and (X_single.at[0,"self_harm_during_admission_Yes"]==0):
    warns.append("PD/PTSD/SUD often present self-harm; both self-harm flags are No — please confirm.")
# ADHD-only 住院少見
_has_adhd_only = (X_single.at[0, "diagnosis_ADHD"]==1) and (sum(X_single.loc[0, [c for c in X_single.columns if c.startswith("diagnosis_")]])==1)
if _has_adhd_only:
    warns.append("ADHD-only inpatient admission is uncommon; check for comorbidities.")
# SCZ/Bipolar LOS 太短
if (X_single.at[0,"diagnosis_Schizophrenia"]==1 or X_single.at[0,"diagnosis_Bipolar"]==1) and (length_of_stay < 7):
    warns.append("Schizophrenia/Bipolar often have longer LOS; LOS<7d looks short — please confirm.")
# 主訴自傷但一開始旗標皆 No → 已自動補 recent self-harm
if ("Suicidal" in complaint or "Self-harm" in complaint) and ("Self-harm" not in " ".join(warns)):
    warns.append("Chief problem indicates self-harm risk — ensured recent self-harm flag = Yes for consistency.")

if warns:
    st.info("ℹ️ Data sanity check:\n- " + "\n- ".join(warns))

# ====== Show results ======
st.subheader("Predicted Dropout Risk (within 3 months)")
c1, c2, c3, c4 = st.columns(4)
with c1: st.metric("Model Probability", f"{percent_model:.1f}%")
with c2: st.metric("Overlay Probability", f"{percent_overlay:.1f}%")
with c3: st.metric("Final Probability", f"{percent_final:.1f}%")
with c4: st.metric("Risk Score (0–100)", f"{score}")

st.caption(
    "Pre-planning: model ignores follow-ups to avoid leakage; Overlay counts planned follow-ups (windowed). "
    "Post-planning: both planned/actual follow-ups are reflected in Overlay. "
    f"Final = (1−{BLEND_W_DEFAULT:.2f})·Model + {BLEND_W_DEFAULT:.2f}·Overlay."
)

if level == "High":
    st.error("🔴 High Risk")
elif level == "Moderate–High":
    st.warning("🟠 Moderate–High (borderline to High)")
elif level == "Moderate":
    st.warning("🟡 Moderate Risk")
elif level == "Low–Moderate":
    st.info("🔵 Low–Moderate (borderline to Moderate)")
else:
    st.success("🟢 Low Risk")

# ====== SHAP + Policy drivers + 對齊卡 ======
with st.expander("🔍 Explanations — Model SHAP vs Policy drivers", expanded=True):
    import xgboost as xgb
    # 單例 SHAP（優先用 pred_contribs）
    try:
        booster = model.get_booster()
        dmat = xgb.DMatrix(X_align, feature_names=list(X_align.columns))
        contribs = booster.predict(dmat, pred_contribs=True, validate_features=False)
        contrib = np.asarray(contribs)[0]
        base_value = float(contrib[-1])
        sv_map = dict(zip(list(X_align.columns), contrib[:-1]))
    except Exception:
        explainer = shap.TreeExplainer(model)
        sv_raw = explainer.shap_values(X_align)
        base_value = explainer.expected_value
        if isinstance(base_value, (list, np.ndarray)) and not np.isscalar(base_value):
            base_value = base_value[0]
            if isinstance(sv_raw, list): sv_raw = sv_raw[0]
        sv_map = dict(zip(list(X_align.columns), sv_raw[0]))

    # 可視化資料
    feat_rows = []
    def _push_shap(label, key, shown_value):
        if key in sv_map:
            feat_rows.append({"feature": label, "value": shown_value, "model_shap": float(sv_map[key]), "key": key})

    _push_shap("Age", "age", X_single.at[0,"age"])
    _push_shap("Length of Stay", "length_of_stay", X_single.at[0,"length_of_stay"])
    _push_shap("Previous Admissions", "num_previous_admissions", X_single.at[0,"num_previous_admissions"])
    _push_shap("Medication Compliance", "medication_compliance_score", X_single.at[0,"medication_compliance_score"])
    _push_shap("Family Support", "family_support_score", X_single.at[0,"family_support_score"])
    _push_shap("Post-discharge Followups", "post_discharge_followups", X_single.at[0,"post_discharge_followups"])
    for dx in diagnoses: _push_shap(f"Diagnosis={dx}", f"diagnosis_{dx}", 1)
    _push_shap(f"Gender={gender}", f"gender_{gender}", 1)
    _push_shap(f"Recent Self-harm={recent_self_harm}", f"has_recent_self_harm_{recent_self_harm}", 1)
    _push_shap(f"Self-harm During Admission={selfharm_adm}", f"self_harm_during_admission_{selfharm_adm}", 1)
    df_shap = pd.DataFrame(feat_rows)

    # SHAP waterfall（取前 12）
    if len(df_shap):
        df_top = df_shap.reindex(df_shap["model_shap"].abs().sort_values(ascending=False).index).head(12)
        names = df_top["feature"].tolist()
        vals = df_top["model_shap"].to_numpy(dtype=float)
        data_vals = df_top["value"].to_numpy(dtype=float)
        exp = shap.Explanation(values=vals, base_values=base_value, feature_names=names, data=data_vals)
        shap.plots.waterfall(exp, show=False, max_display=12)
        st.pyplot(plt.gcf(), clear_figure=True)
    else:
        st.caption("No SHAP contributions available for the selected case.")

    # SHAP 表格
    st.caption("Model SHAP (top by |value|)")
    if len(df_shap):
        st.dataframe(
            df_shap.reindex(df_shap["model_shap"].abs().sort_values(ascending=False).index)[["feature","value","model_shap"]].head(12),
            use_container_width=True
        )

    # Policy drivers
    df_drv = pd.DataFrame(
        [{"driver": k, "policy_log_odds (pre-scale)": round(v, 3)} for k, v in sorted(drivers, key=lambda x: abs(x[1]), reverse=True)]
    )
    st.caption("Policy drivers")
    if len(df_drv):
        st.dataframe(df_drv, use_container_width=True)
    else:
        st.write("No policy drivers for this case.")

    # Model vs Policy 對齊
    st.caption("Alignment check (Model vs Policy) — look for ⚠️ if directions disagree.")
    def _sign(x): return 1 if x>1e-6 else (-1 if x<-1e-6 else 0)
    align_rows = []
    name_map = [
        ("Previous Admissions","num_previous_admissions","More previous admissions"),
        ("Medication Compliance","medication_compliance_score","Low medication compliance"),
        ("Family Support","family_support_score","Low family support"),
        ("Followups","post_discharge_followups","Follow-ups within"),
        ("Length of Stay","length_of_stay","Very short stay"),
    ]
    for lab, key, dname in name_map:
        shap_v = float(sv_map.get(key, 0.0))
        pol = 0.0
        for nm, v in drivers:
            if dname.split()[0] in nm: pol += v
        ms, ps = _sign(shap_v), _sign(pol)
        align_rows.append({"feature": lab, "model_sign": ms, "policy_sign": ps, "flag": "⚠️" if (ms*ps==-1) else ""})
    st.dataframe(pd.DataFrame(align_rows), use_container_width=True)

# ====== Recommended actions（簡化版） ======
st.subheader("Recommended Actions")
BASE_ACTIONS = {
    "High": [
        ("Today","Clinician","Crisis/safety planning; 24/7 crisis contacts"),
        ("Today","Clinic scheduler","Return within 7 days (prefer 72h)"),
        ("Today","Care coordinator","Warm handoff to case management"),
        ("48h","Nurse","Outreach call: symptoms/side-effects/barriers"),
        ("7d","Pharmacist/Nurse","Medication review + adherence aids"),
        ("1–2w","Social worker","SDOH screen; transport/financial aid"),
        ("1–4w","Peer support","Enroll in peer/skills group"),
    ],
    "Moderate": [
        ("1–2w","Clinic scheduler","Book within 14 days; SMS reminders"),
        ("1–2w","Nurse","Barrier check & solutions"),
        ("2–4w","Clinician","Brief MI/BA/psychoeducation; 4-week plan"),
    ],
    "Low": [
        ("2–4w","Clinic scheduler","Routine follow-up; confirm reminders"),
        ("2–4w","Nurse","Education/self-management resources"),
    ],
}
def _normalize_action_tuple(a):
    if len(a)==3: return (a[0],a[1],a[2])
    return a

def personalized_actions(row: pd.Series, chosen_dx: list):
    acts = []
    comp = _num_or_default(row["medication_compliance_score"], "medication_compliance_score")
    sup  = _num_or_default(row["family_support_score"], "family_support_score")
    fup  = _num_or_default(row["post_discharge_followups"], "post_discharge_followups")
    los  = _num_or_default(row["length_of_stay"], "length_of_stay")
    agev = _num_or_default(row["age"], "age")
    has_selfharm = flag_yes(row, "has_recent_self_harm") or flag_yes(row, "self_harm_during_admission")
    has_sud = "Substance Use Disorder" in chosen_dx
    has_pd  = "Personality Disorder" in chosen_dx
    has_dep = "Depression" in chosen_dx
    has_scz = "Schizophrenia" in chosen_dx

    if has_selfharm:
        acts += [("Today","Clinician","C-SSRS; update safety plan; lethal-means counseling"),
                 ("48h","Nurse","Safety check-in call")]
    if has_sud and comp <= 3:
        acts += [("1–7d","Clinician","Brief MI focused on use goals"),
                 ("1–7d","Care coordinator","Refer to SUD program/IOP or CM"),
                 ("Today","Clinician","Overdose prevention education")]
    if has_pd and los < 3:
        acts += [("Today","Care coordinator","Same-day DBT/skills intake"),
                 ("48h","Peer support","Proactive outreach + skills workbook")]
    if comp <= 3:
        acts += [("7d","Pharmacist","Simplify regimen + blister/pillbox + reminders"),
                 ("1–2w","Clinician","Consider LAI if appropriate")]
    if sup <= 2:
        acts += [("1–2w","Clinician","Family meeting / caregiver engagement"),
                 ("1–2w","Social worker","Community supports; transport/financial counseling")]
    if fup == 0:
        acts += [("Today","Clinic scheduler","Book 2 touchpoints in first 14 days (day2/day7)")]
    if los < 3:
        acts += [("48h","Nurse","Early call; review meds/barriers")]
    elif los > 21:
        acts += [("1–7d","Care coordinator","Step-down/day program plan + warm handoff")]
    if agev < 21:
        acts += [("1–2w","Clinician","Involve guardians; link school counseling")]
    elif agev >= 75:
        acts += [("1–2w","Nurse/Pharmacist","Med reconciliation; simplify dosing")]
    if has_dep:
        acts += [("1–2w","Clinician","Behavioral activation + activity schedule")]
    if has_scz:
        acts += [("1–4w","Clinician","Relapse plan; early warning signs; caregiver involvement")]
    return acts

bucket = {"High":"High","Moderate–High":"High","Moderate":"Moderate","Low–Moderate":"Low","Low":"Low"}
acts = [ _normalize_action_tuple(a) for a in BASE_ACTIONS[bucket[level]] ]
acts += personalized_actions(X_single.iloc[0], diagnoses)
seen=set(); uniq=[]
for a in acts:
    key=(a[0],a[1],a[2])
    if key not in seen:
        seen.add(key); uniq.append(a)
ORDER={"Today":0,"48h":1,"7d":2,"1–7d":2,"1–2w":3,"2–4w":4,"1–4w":5}
uniq.sort(key=lambda x:(ORDER.get(x[0],99), x[1], x[2]))
st.dataframe(pd.DataFrame(uniq, columns=["Timeline","Owner","Action"]), use_container_width=True)

if level in ["High","Moderate–High"]:
    def make_sop_txt(score, label, actions):
        lines=["Psychiatric Dropout Risk – SOP", f"Risk score: {score}/100 | Risk level: {label}",""]
        for (tl,ow,ac) in actions: lines.append(f"- {tl} | {ow} | {ac}")
        buf=BytesIO("\n".join(lines).encode("utf-8")); buf.seek(0); return buf
    st.download_button("⬇️ Export SOP (TXT)", make_sop_txt(score, level, uniq),
                       file_name="dropout_risk_SOP.txt", mime="text/plain")
    if HAS_DOCX:
        def make_docx(score,label,actions):
            doc=Document(); doc.add_heading('Psychiatric Dropout Risk – SOP', level=1)
            doc.add_paragraph(f"Risk score: {score}/100 | Risk level: {label}")
            t=doc.add_table(rows=1, cols=3); hdr=t.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = 'Timeline','Owner','Action'
            for (tl,ow,ac) in actions:
                r=t.add_row().cells; r[0].text, r[1].text, r[2].text = tl,ow,ac
            buf=BytesIO(); doc.save(buf); buf.seek(0); return buf
        st.download_button("⬇️ Export SOP (Word)", make_docx(score, level, uniq),
                           file_name="dropout_risk_SOP.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ====== What-if 小面板 ======
with st.expander("🧪 What-if: adjust followups/compliance and recompute", expanded=False):
    wf_follow = st.slider("What-if followups (count)", 0, 10, int(X_single.at[0,"post_discharge_followups"]))
    wf_comp   = st.slider("What-if compliance", 0.0, 10.0, float(X_single.at[0,"medication_compliance_score"]), 0.5)
    wf_window = st.number_input("What-if follow-up window (days)", min_value=3, max_value=60, value=int(followup_window_days), step=1)
    X_wf = X_single.copy()
    X_wf.at[0,"post_discharge_followups"] = wf_follow
    X_wf.at[0,"medication_compliance_score"] = wf_comp
    X_wf_al, _ = align_df_to_model(X_wf, model)
    p_m_wf = float(predict_model_proba(X_wf_al)[0]) if use_followups_feature else float(predict_model_proba(align_df_to_model(X_wf.assign(post_discharge_followups=0), model)[0])[0])
    p_o_wf, _ = overlay_single_and_drivers(X_wf, followup_window_days=int(wf_window))
    p_f_wf = (1.0 - BLEND_W_DEFAULT) * p_m_wf + BLEND_W_DEFAULT * p_o_wf
    st.write(f"Model={p_m_wf*100:.1f}% | Overlay={p_o_wf*100:.1f}% | Final={p_f_wf*100:.1f}%")

# ====== Batch Prediction ======
st.markdown("---")
st.subheader("Batch Prediction (Excel)")
friendly_cols = [
    "Age","Gender","Diagnoses","Length of Stay (days)","Previous Admissions (1y)",
    "Medication Compliance Score (0–10)","Family Support Score (0–10)",
    "Post-discharge Followups (count)","Follow-up Window (days)",
    "Recent Self-harm","Self-harm During Admission","Chief problem"
]
tpl = pd.DataFrame(columns=friendly_cols)
buf_tpl = BytesIO(); tpl.to_excel(buf_tpl, index=False); buf_tpl.seek(0)
st.download_button("📥 Download Excel Template", buf_tpl, file_name="batch_template.xlsx",
                   mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

uploaded = st.file_uploader("📂 Upload Excel", type=["xlsx"])
def parse_multi(cell):
    parts = [p.strip() for p in re.split(r"[;,/|]", str(cell)) if p.strip()]
    return parts if parts else []

if uploaded is not None:
    try:
        raw = pd.read_excel(uploaded)
        df = pd.DataFrame(0, index=raw.index, columns=TEMPLATE_COLUMNS, dtype=float)

        def safe(col, default=0): return raw[col] if col in raw.columns else default
        # 連續
        map_cont = [
            ("Age","age"), ("Length of Stay (days)","length_of_stay"),
            ("Previous Admissions (1y)","num_previous_admissions"),
            ("Medication Compliance Score (0–10)","medication_compliance_score"),
            ("Family Support Score (0–10)","family_support_score"),
        ]
        for k_raw, k in map_cont:
            df[k] = safe(k_raw)

        # followups 數量
        if "Post-discharge Followups (count)" in raw.columns:
            df["post_discharge_followups"] = raw["Post-discharge Followups (count)"]
        else:
            df["post_discharge_followups"] = 0

        # one-hots
        if "Gender" in raw.columns:
            for i, v in raw["Gender"].astype(str).str.strip().items():
                col=f"gender_{v}"
                if col in df.columns: df.at[i,col]=1
        if "Diagnoses" in raw.columns:
            for i, cell in raw["Diagnoses"].items():
                for v in parse_multi(cell):
                    col=f"diagnosis_{v}"
                    if col in df.columns: df.at[i,col]=1
        for col_h, pre in [("Recent Self-harm","has_recent_self_harm"), ("Self-harm During Admission","self_harm_during_admission")]:
            if col_h in raw.columns:
                for i, v in raw[col_h].astype(str).str.strip().items():
                    col=f"{pre}_{v}"
                    if col in df.columns: df.at[i,col]=1

        # Chief problem 一致化：Suicidal/Self-harm 主訴→若兩旗標皆 No，補 recent self harm = Yes
        if "Chief problem" in raw.columns:
            for i, cp in raw["Chief problem"].astype(str).items():
                if ("Suicidal" in cp) or ("Self-harm" in cp):
                    if (df.at[i,"has_recent_self_harm_Yes"]==0) and (df.at[i,"self_harm_during_admission_Yes"]==0):
                        df.at[i,"has_recent_self_harm_Yes"]=1; df.at[i,"has_recent_self_harm_No"]=0

        fill_defaults_batch(df)

        # Pre-planning → 模型不吃 followups
        if not use_followups_feature:
            df["post_discharge_followups"] = 0

        Xb_al, _ = align_df_to_model(df, model)
        base_probs = predict_model_proba(Xb_al)

        # Overlay vectorized（含 follow-up window 權重）
        def overlay_vec(df_feat: pd.DataFrame, base_probs, followup_windows=None):
            base = _logit_vec(base_probs); lz = base.copy()
            adm = pd.to_numeric(df_feat["num_previous_admissions"], errors="coerce").fillna(DEFAULTS["num_previous_admissions"]).to_numpy()
            comp= pd.to_numeric(df_feat["medication_compliance_score"], errors="coerce").fillna(DEFAULTS["medication_compliance_score"]).to_numpy()
            sup = pd.to_numeric(df_feat["family_support_score"], errors="coerce").fillna(DEFAULTS["family_support_score"]).to_numpy()
            fup = pd.to_numeric(df_feat["post_discharge_followups"], errors="coerce").fillna(DEFAULTS["post_discharge_followups"]).to_numpy()
            los = pd.to_numeric(df_feat["length_of_stay"], errors="coerce").fillna(DEFAULTS["length_of_stay"]).to_numpy()
            agev= pd.to_numeric(df_feat["age"], errors="coerce").fillna(DEFAULTS["age"]).to_numpy()

            lz += POLICY["per_prev_admission"] * np.minimum(adm, 5)
            lz += POLICY["per_point_low_support"] * np.maximum(0.0, 5.0 - sup)
            lz += POLICY["per_point_low_compliance"] * np.maximum(0.0, 5.0 - comp)
            lz += POLICY["per_point_high_compliance_protect"] * np.maximum(0.0, comp - 7.0)

            # window weights
            if followup_windows is None:
                wnd = np.full_like(fup, 14.0, dtype=float)
            else:
                wnd = pd.to_numeric(followup_windows, errors="coerce").fillna(14.0).to_numpy()
                wnd = np.clip(wnd, 3, 60)
            w = np.sqrt(14.0 / wnd)
            w = np.clip(w, 0.5, 1.5)
            eff_followups = fup * w
            lz += POLICY["per_followup"] * eff_followups
            lz += POLICY["no_followup_extra"] * (fup == 0)

            lz += np.where(los < 3, POLICY["los_short"],
                    np.where(los <= 14, POLICY["los_mid"],
                    np.where(los <= 21, POLICY["los_mid_high"], POLICY["los_long"])))
            lz += POLICY["age_young"] * (agev < 21) + POLICY["age_old"] * (agev >= 75)

            for dx, w in POLICY["diag"].items():
                col=f"diagnosis_{dx}"
                if col in df_feat.columns: lz += w * (df_feat[col].to_numpy() == 1)
            sud = (df_feat.get("diagnosis_Substance Use Disorder",0).to_numpy()==1)
            pdm = (df_feat.get("diagnosis_Personality Disorder",0).to_numpy()==1)
            lz += POLICY["x_sud_lowcomp"] * (sud & (comp <= 3))
            lz += POLICY["x_pd_shortlos"] * (pdm & (los < 3))

            delta = np.clip(OVERLAY_SCALE * (lz - base), -DELTA_CLIP, DELTA_CLIP)
            lz2 = base + delta + CAL_LOGIT_SHIFT
            return 1.0 / (1.0 + np.exp(-(lz2 / TEMP)))

        # 取批次 window 欄位
        fw_col = raw["Follow-up Window (days)"] if "Follow-up Window (days)" in raw.columns else None
        p_overlay_b = overlay_vec(df, base_probs, followup_windows=fw_col)
        p_final_b = (1.0 - BLEND_W_DEFAULT) * base_probs + BLEND_W_DEFAULT * p_overlay_b

        # 自傷 uplift
        hr = df.get("has_recent_self_harm_Yes", 0); ha = df.get("self_harm_during_admission_Yes", 0)
        mask = ((np.array(hr)==1) | (np.array(ha)==1))
        p_final_b[mask] = np.minimum(np.maximum(p_final_b[mask], SOFT_UPLIFT["floor"]) + SOFT_UPLIFT["add"], SOFT_UPLIFT["cap"])

        out = raw.copy()
        out["risk_percent"] = (p_final_b*100).round(1)
        out["risk_score_0_100"] = (p_final_b*100).round().astype(int)

        s = out["risk_score_0_100"].to_numpy()
        levels = np.full(s.shape, "Low", dtype=object)
        levels[s >= 20 - BORDER_BAND] = "Low–Moderate"
        levels[s >= 20 + BORDER_BAND] = "Moderate"
        levels[s >= 40 - BORDER_BAND] = "Moderate–High"
        levels[s >= 40 + BORDER_BAND] = "High"
        out["risk_level"] = levels

        st.dataframe(out, use_container_width=True)
        buf = BytesIO(); out.to_csv(buf, index=False); buf.seek(0)
        st.download_button("⬇️ Download Results (CSV)", buf, "predictions.csv", "text/csv")
    except Exception as e:
        st.error(f"Batch error: {e}")

# ====== Validation（Ablation + Decision curve + Capacity + Fairness）======
st.markdown("---")
st.header("✅ Validation (synthetic hold-out)")

cA, cB = st.columns([2,1])
with cA:
    n_val = st.slider("Number of synthetic patients", 5000, 60000, 20000, 5000)
    seed_val = st.number_input("Random seed", min_value=1, max_value=10**9, value=2024, step=1)
with cB:
    pt_low = st.slider("Decision threshold (Moderate%):", 5, 60, 20, 1)
    pt_high = st.slider("Decision threshold (High%):", 20, 90, 40, 1)
run_val = st.button("Run validation")

def generate_synth_holdout(n=20000, seed=2024):
    rng = np.random.default_rng(seed)
    df = pd.DataFrame(0, index=range(n), columns=TEMPLATE_COLUMNS, dtype=float)
    df["age"] = rng.integers(16, 85, n)

    los = rng.normal(21.0, 7.0, n)
    los = np.clip(los, 1, 60)
    idx_primary = rng.integers(0, len(DIAG_LIST), n)
    for i, d in enumerate(DIAG_LIST): df.loc[idx_primary == i, f"diagnosis_{d}"] = 1
    scz = (df["diagnosis_Schizophrenia"] == 1).to_numpy()
    bp  = (df["diagnosis_Bipolar"] == 1).to_numpy()
    los = los + 7.0*scz + 4.0*bp
    df["length_of_stay"] = np.clip(los, 1, 90)

    df["num_previous_admissions"] = rng.poisson(0.9, n).clip(0, 12)
    df["medication_compliance_score"] = rng.normal(6.0, 2.5, n).clip(0, 10)
    df["family_support_score"] = rng.normal(5.0, 2.5, n).clip(0, 10)
    df["post_discharge_followups"] = rng.integers(0, 6, n)
    idx_gender = rng.integers(0, len(GENDER_LIST), n)
    for i, g in enumerate(GENDER_LIST): df.loc[idx_gender == i, f"gender_{g}"] = 1
    extra_probs = {"Substance Use Disorder": 0.20, "Anxiety": 0.20, "Depression": 0.25, "PTSD": 0.10}
    for d, pr in extra_probs.items(): df.loc[rng.random(n) < pr, f"diagnosis_{d}"] = 1

    # 自傷旗標分布
    base_r = rng.random(n)
    pd_flag   = (df["diagnosis_Personality Disorder"] == 1).to_numpy()
    ptsd_flag = (df["diagnosis_PTSD"] == 1).to_numpy()
    sud_flag  = (df["diagnosis_Substance Use Disorder"] == 1).to_numpy()
    prob_sh = 0.18 + 0.12*pd_flag + 0.10*ptsd_flag + 0.10*sud_flag
    r1 = (base_r < np.clip(prob_sh, 0, 0.9)).astype(int)
    r2 = (rng.random(n) < 0.12).astype(int)
    df.loc[r1 == 1, "has_recent_self_harm_Yes"] = 1; df.loc[r1 == 0, "has_recent_self_harm_No"] = 1
    df.loc[r2 == 1, "self_harm_during_admission_Yes"] = 1; df.loc[r2 == 0, "self_harm_during_admission_No"] = 1

    # 真值（與訓練合成一致）
    beta0 = -0.60
    beta = {"has_recent_self_harm_Yes":0.80,"self_harm_during_admission_Yes":0.60,
            "prev_adm_ge2":0.60,"medication_compliance_per_point":-0.25,"family_support_per_point":-0.20,
            "followups_per_visit":-0.15,"length_of_stay_per_day":0.04}
    beta_diag = {"Personality Disorder":0.35,"Substance Use Disorder":0.35,"Bipolar":0.10,"PTSD":0.10,"Schizophrenia":0.10,"Depression":0.05,
                 "Anxiety":0.00,"OCD":0.00,"Dementia":0.00,"ADHD":0.00,"Other/Unknown":0.00}
    prev_ge2 = (df["num_previous_admissions"] >= 2).astype(np.float32)
    logit = (beta0
             + beta["has_recent_self_harm_Yes"]*df["has_recent_self_harm_Yes"]
             + beta["self_harm_during_admission_Yes"]*df["self_harm_during_admission_Yes"]
             + beta["prev_adm_ge2"]*prev_ge2
             + beta["medication_compliance_per_point"]*df["medication_compliance_score"]
             + beta["family_support_per_point"]*df["family_support_score"]
             + beta["followups_per_visit"]*df["post_discharge_followups"]
             + beta["length_of_stay_per_day"]*df["length_of_stay"])
    for d,w in beta_diag.items(): logit = logit + w*df[f"diagnosis_{d}"]
    noise = np.random.default_rng(seed+1).normal(0.0, 0.35, n).astype(np.float32)
    p_true = 1.0 / (1.0 + np.exp(-(logit + noise)))
    y_true = (np.random.default_rng(seed+2).random(n) < p_true).astype(int)

    fill_defaults_batch(df)
    return df, y_true

def plot_roc_pr(y, p_list, labels):
    from sklearn.metrics import roc_curve, auc, precision_recall_curve, average_precision_score
    # ROC
    fig1, ax1 = plt.subplots()
    for p,l in zip(p_list, labels):
        fpr, tpr, _ = roc_curve(y, p); roc = auc(fpr, tpr)
        ax1.plot(fpr, tpr, label=f"{l} AUC={roc:.3f}")
    ax1.plot([0,1],[0,1],"--")
    ax1.set_xlabel("FPR"); ax1.set_ylabel("TPR"); ax1.set_title("ROC")
    ax1.legend(loc="lower right"); st.pyplot(fig1, clear_figure=True)
    # PR
    fig2, ax2 = plt.subplots()
    for p,l in zip(p_list, labels):
        prec, rec, _ = precision_recall_curve(y, p); ap = average_precision_score(y, p)
        ax2.plot(rec, prec, label=f"{l} AP={ap:.3f}")
    ax2.set_xlabel("Recall"); ax2.set_ylabel("Precision"); ax2.set_title("PR")
    ax2.legend(loc="upper right"); st.pyplot(fig2, clear_figure=True)

def ece(y, p, n_bins=10):
    bins = np.linspace(0.0,1.0,n_bins+1)
    idx = np.digitize(p, bins)-1
    err=0.0
    for b in range(n_bins):
        m=(idx==b)
        if m.sum()==0: continue
        fp=p[m].mean(); tp=y[m].mean()
        err += m.mean()*abs(tp-fp)
    return float(err)

def confusion(y, p, thr):
    from sklearn.metrics import confusion_matrix
    yhat = (p>=thr).astype(int)
    tn,fp,fn,tp = confusion_matrix(y,yhat).ravel()
    return tn,fp,fn,tp

def decision_curve(y, p):
    # Net Benefit across thresholds 0.05~0.60
    ths = np.linspace(0.05,0.60,56)
    N=len(y)
    nb=[]
    for t in ths:
        yhat = (p>=t).astype(int)
        tp = ((yhat==1)&(y==1)).sum(); fp=((yhat==1)&(y==0)).sum()
        nb.append((tp/N) - (fp/N)*(t/(1-t)))
    return ths, np.array(nb)

if run_val:
    try:
        from sklearn.metrics import roc_auc_score, average_precision_score, brier_score_loss
    except Exception as e:
        st.error(f"Need scikit-learn: {e}")
    else:
        with st.spinner("Generating data & evaluating..."):
            df_syn, y_true = generate_synth_holdout(n_val, seed_val)
            if not use_followups_feature:
                df_syn["post_discharge_followups"] = 0
            Xa, _ = align_df_to_model(df_syn, model)
            p_model_v = predict_model_proba(Xa)

            # Overlay-only（BLEND=1）
            def overlay_only_vec(df_feat, base_probs):
                base = _logit_vec(base_probs); lz = base.copy()
                adm = pd.to_numeric(df_feat["num_previous_admissions"], errors="coerce").fillna(DEFAULTS["num_previous_admissions"]).to_numpy()
                comp= pd.to_numeric(df_feat["medication_compliance_score"], errors="coerce").fillna(DEFAULTS["medication_compliance_score"]).to_numpy()
                sup = pd.to_numeric(df_feat["family_support_score"], errors="coerce").fillna(DEFAULTS["family_support_score"]).to_numpy()
                fup = pd.to_numeric(df_feat["post_discharge_followups"], errors="coerce").fillna(DEFAULTS["post_discharge_followups"]).to_numpy()
                los = pd.to_numeric(df_feat["length_of_stay"], errors="coerce").fillna(DEFAULTS["length_of_stay"]).to_numpy()
                agev= pd.to_numeric(df_feat["age"], errors="coerce").fillna(DEFAULTS["age"]).to_numpy()

                lz += POLICY["per_prev_admission"] * np.minimum(adm, 5)
                lz += POLICY["per_point_low_support"] * np.maximum(0.0, 5.0 - sup)
                lz += POLICY["per_point_low_compliance"] * np.maximum(0.0, 5.0 - comp)
                lz += POLICY["per_point_high_compliance_protect"] * np.maximum(0.0, comp - 7.0)

                # Overlay window：驗證時固定 14d 權重（或可做成 slider）
                w = np.sqrt(14.0 / 14.0)
                eff_followups = fup * np.clip(w, 0.5, 1.5)
                lz += POLICY["per_followup"] * eff_followups
                lz += POLICY["no_followup_extra"] * (fup == 0)

                lz += np.where(los < 3, POLICY["los_short"],
                        np.where(los <= 14, POLICY["los_mid"],
                        np.where(los <= 21, POLICY["los_mid_high"], POLICY["los_long"])))
                lz += POLICY["age_young"] * (agev < 21) + POLICY["age_old"]*(agev >= 75)
                for dx,wg in POLICY["diag"].items():
                    col=f"diagnosis_{dx}"
                    if col in df_feat.columns: lz += wg * (df_feat[col].to_numpy()==1)
                sud = (df_feat.get("diagnosis_Substance Use Disorder",0).to_numpy()==1)
                pdm = (df_feat.get("diagnosis_Personality Disorder",0).to_numpy()==1)
                lz += POLICY["x_sud_lowcomp"] * (sud & (comp <= 3))
                lz += POLICY["x_pd_shortlos"] * (pdm & (los < 3))
                delta = np.clip(OVERLAY_SCALE * (lz - base), -DELTA_CLIP, DELTA_CLIP)
                lz2 = base + delta + CAL_LOGIT_SHIFT
                return 1.0 / (1.0 + np.exp(-(lz2 / TEMP)))

            p_overlay_v = overlay_only_vec(df_syn, p_model_v)
            p_final_v = (1.0 - BLEND_W_DEFAULT) * p_model_v + BLEND_W_DEFAULT * p_overlay_v

            # Metrics 表
            auc_m = roc_auc_score(y_true, p_model_v); auc_o = roc_auc_score(y_true, p_overlay_v); auc_f = roc_auc_score(y_true, p_final_v)
            ap_m  = average_precision_score(y_true, p_model_v); ap_o  = average_precision_score(y_true, p_overlay_v); ap_f  = average_precision_score(y_true, p_final_v)
            br_m  = brier_score_loss(y_true, p_model_v); br_o  = brier_score_loss(y_true, p_overlay_v); br_f  = brier_score_loss(y_true, p_final_v)
            ece_m = ece(y_true, p_model_v); ece_o = ece(y_true, p_overlay_v); ece_f = ece(y_true, p_final_v)
            st.subheader("Ablation metrics")
            st.dataframe(pd.DataFrame([
                {"Model":"Model","AUC":auc_m,"PR-AUC":ap_m,"Brier":br_m,"ECE(10bins)":ece_m},
                {"Model":"Overlay-only","AUC":auc_o,"PR-AUC":ap_o,"Brier":br_o,"ECE(10bins)":ece_o},
                {"Model":"Blend (Final)","AUC":auc_f,"PR-AUC":ap_f,"Brier":br_f,"ECE(10bins)":ece_f},
            ]).round(3), use_container_width=True)

            st.subheader("Curves")
            plot_roc_pr(y_true, [p_model_v, p_overlay_v, p_final_v], ["Model","Overlay","Final"])

            # Decision Curve
            ths_m, nb_m = decision_curve(y_true, p_model_v)
            ths_o, nb_o = decision_curve(y_true, p_overlay_v)
            ths_f, nb_f = decision_curve(y_true, p_final_v)
            fig, ax = plt.subplots()
            ax.plot(ths_m, nb_m, label="Model")
            ax.plot(ths_o, nb_o, label="Overlay")
            ax.plot(ths_f, nb_f, label="Final")
            ax.axhline(0, ls="--"); ax.set_xlabel("Threshold"); ax.set_ylabel("Net benefit"); ax.set_title("Decision Curve")
            ax.legend(); st.pyplot(fig, clear_figure=True)

            # Confusion + Capacity（以門檻計）
            thr_mod = pt_low/100.0; thr_hi = pt_high/100.0
            st.subheader("Operational (binary @ thresholds)")
            c1, c2 = st.columns(2)
            with c1:
                st.markdown(f"**Moderate ≥{pt_low}%**")
                for name, p in [("Model",p_model_v),("Overlay",p_overlay_v),("Final",p_final_v)]:
                    tn,fp,fn,tp = confusion(y_true, p, thr_mod)
                    st.write(f"{name} — TN:{tn} FP:{fp} FN:{fn} TP:{tp}")
            with c2:
                st.markdown(f"**High ≥{pt_high}%**")
                for name, p in [("Model",p_model_v),("Overlay",p_overlay_v),("Final",p_final_v)]:
                    tn,fp,fn,tp = confusion(y_true, p, thr_hi)
                    st.write(f"{name} — TN:{tn} FP:{fp} FN:{fn} TP:{tp}")

            # Capacity & time load（簡易估算）
            st.subheader("Capacity / Time load (per 1,000 patients)")
            N = 1000
            def count_at(p, thr): return int(((p>=thr).sum() / len(p)) * N)
            mod_cnt = count_at(p_final_v, thr_mod); high_cnt = count_at(p_final_v, thr_hi)
            st.caption("Assumptions (editable):")
            t_outreach = st.number_input("Nurse outreach (min)", 5, 60, 15, 5)
            t_sched = st.number_input("Scheduler booking (min)", 2, 30, 5, 1)
            t_pharm = st.number_input("Pharmacist review (min)", 5, 60, 20, 5)
            hours = (mod_cnt*(t_outreach+t_sched) + (high_cnt)*(t_pharm)) / 60.0
            st.write(f"Flagged Moderate+ per 1,000: **{mod_cnt}** ; High: **{high_cnt}** → ~ **{hours:.1f} hours** total effort.")

            # Fairness：性別/年齡段/主診斷
            st.subheader("Fairness (AUC / ECE by subgroup, Final)")
            def subgroup_auc_ece(df_feat, y, p, mask):
                if mask.sum()<100: return np.nan, np.nan
                from sklearn.metrics import roc_auc_score
                return float(roc_auc_score(y[mask], p[mask])), float(ece(y[mask], p[mask], 10))
            agev = pd.to_numeric(df_syn["age"], errors="coerce").fillna(40).to_numpy()
            bands = {"<30": (agev<30), "30–59": ((agev>=30)&(agev<60)), "≥60": (agev>=60)}
            rows=[]
            for g in GENDER_LIST:
                mask = (df_syn.get(f"gender_{g}",0).to_numpy()==1)
                a,e = subgroup_auc_ece(df_syn, y_true, p_final_v, mask)
                rows.append({"group":"Gender", "value":g, "AUC":a, "ECE":e})
            for k,m in bands.items():
                a,e = subgroup_auc_ece(df_syn, y_true, p_final_v, m)
                rows.append({"group":"AgeBand", "value":k, "AUC":a, "ECE":e})
            diag_cols=[f"diagnosis_{d}" for d in DIAG_LIST]
            prim = np.argmax(df_syn[diag_cols].to_numpy(), axis=1)
            for i,d in enumerate(DIAG_LIST):
                mask = (prim==i)
                a,e = subgroup_auc_ece(df_syn, y_true, p_final_v, mask)
                rows.append({"group":"PrimaryDx", "value":d, "AUC":a, "ECE":e})
            st.dataframe(pd.DataFrame(rows).round(3), use_container_width=True)

        st.success("Validation finished.")

# ====== Vignettes（for expert review）======
st.markdown("---")
st.header("🧾 Vignettes template (for expert review)")
def _mk_vignette_row(age, gender, diags, los, prev, comp, rsh, shadm, sup, fup):
    return {
        "Age": age, "Gender": gender, "Diagnoses": ", ".join(diags),
        "Length of Stay (days)": los, "Previous Admissions (1y)": prev,
        "Medication Compliance Score (0–10)": comp,
        "Family Support Score (0–10)": sup,
        "Post-discharge Followups (count)": fup,
        "Recent Self-harm": rsh, "Self-harm During Admission": shadm,
        "Expert Risk (Low/Moderate/High or 0–100)": ""
    }
def build_vignettes_df(n=20, seed=77):
    rng = np.random.default_rng(seed); base=[]
    protos = [
        (19,"Female",["Depression"],2,0,3,"No","No",2,0),
        (28,"Male",["Substance Use Disorder"],1,3,2,"No","No",3,0),
        (35,"Male",["Bipolar"],6,1,4,"No","No",5,1),
        (42,"Female",["Personality Disorder"],2,2,3,"No","No",4,0),
        (55,"Male",["Schizophrenia"],10,4,5,"No","No",5,2),
        (63,"Female",["PTSD"],4,1,6,"No","No",6,1),
        (72,"Male",["Depression","Anxiety"],5,0,7,"No","No",7,2),
        (23,"Female",["OCD"],3,0,8,"No","No",8,2),
        (31,"Male",["Substance Use Disorder","Depression"],7,2,3,"No","No",3,0),
        (47,"Female",["Personality Disorder","PTSD"],2,1,4,"No","No",4,0),
        (38,"Male",["ADHD"],4,0,6,"No","No",6,1),
        (26,"Female",["Anxiety"],1,0,5,"No","No",5,1),
        (60,"Male",["Dementia"],12,1,6,"No","No",6,2),
        (45,"Female",["Schizophrenia","Substance Use Disorder"],9,3,2,"No","No",3,0),
        (52,"Male",["Bipolar","Personality Disorder"],2,2,3,"No","No",4,0),
        (33,"Female",["Depression"],3,0,8,"No","No",8,3),
        (29,"Male",["Substance Use Disorder"],5,2,2,"No","No",4,0),
        (70,"Female",["Depression","PTSD"],8,1,5,"No","No",6,2),
        (41,"Male",["Personality Disorder","Substance Use Disorder"],2,3,3,"No","No",3,0),
        (36,"Female",["Other/Unknown"],4,0,5,"No","No",5,1),
    ]
    for i in range(min(n, len(protos))): base.append(_mk_vignette_row(*protos[i]))
    for _ in range(len(base), n):
        age = int(np.clip(rng.normal(40,15), 18, 90))
        gender = GENDER_LIST[int(rng.integers(0,len(GENDER_LIST)))]
        k = int(rng.integers(1,3)); diags = list(rng.choice(DIAG_LIST, size=k, replace=False))
        los = int(np.clip(rng.normal(21,7),1,60)); prev = int(np.clip(rng.poisson(1.0),0,8))
        comp = float(np.clip(rng.normal(6,2.5),0,10)); rsh = rng.choice(["Yes","No"]); shadm = rng.choice(["Yes","No"])
        sup = float(np.clip(rng.normal(5,2.5),0,10)); fup = int(np.clip(rng.integers(0,4),0,10))
        base.append(_mk_vignette_row(age, gender, diags, los, prev, comp, rsh, shadm, sup, fup))
    return pd.DataFrame(base)

vdf = build_vignettes_df(20, 77)
buf_v = BytesIO(); vdf.to_excel(buf_v, index=False); buf_v.seek(0)
st.download_button("📥 Download Vignettes (20 cases, Excel)", buf_v,
                   file_name="vignettes_20.xlsx",
                   mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# ====== Data dictionary ======
with st.expander("📚 Data dictionary / Definitions", expanded=False):
    st.markdown(f"""
- **Medication Compliance (0–10)**：0=幾乎不服藥；10=幾乎完全依從（近 1 個月）
- **Family Support (0–10)**：0=非常不足；10=非常充足
- **Post-discharge Followups (count)**：出院後 **Follow-up window (days)** 內的「門診回診 / 電話關懷 / 社工或個管接觸」次數總和。
- **Follow-up window (days)**：計算 followups 的時間窗；**視窗越短，保護效果越強**（以 14 天為基準做權重）。
- **Pre-planning**：模型不吃 followups（避免洩漏），但 **Overlay** 仍將「已規劃的追蹤」視為保護，**Final** 因而改變。
- **Post-planning**：Overlay 反映實際追蹤帶來的風險變化，便於動態監測。
- **Final Probability**：Model 與 Policy Overlay 的混合（BLEND={BLEND_W_DEFAULT:.2f}），含必要安全 uplift（自傷）。
""")
